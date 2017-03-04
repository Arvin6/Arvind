'''
Created on Mar 4, 2017
Docx module
@author: Arvind
'''
import docx
from package1.formatfactory import formatfactory

class docxwrite(formatfactory):
    '''
    Save as docx
    '''
    def __init__(self):
        '''
        Constructor
        '''
        
    def _write(self,details,saveloc):
            mov = docx.Document()
            mov.add_paragraph("Name of the Movie:\t"+details['name'])
            mov.add_paragraph("Running time:\t\t"+details['time'])
            mov.add_paragraph("Language:\t\t"+details['language'])
            mov.add_paragraph("Lead actor:\t\t"+details['act'])
            mov.add_paragraph("Genre:\t\t\t"+details['genre'])
            mov.save(saveloc+"/Movie.docx")
            #mov.save('G://Eclipse workspace/Movie-Aggregator/package1/Result/Movie.docx')     
        
if __name__ == '__main__':                              
    '''For testing'''
    det={'genre': 'Action', 'act': 'Hugh Jackman', 'name': 'Logan', 'language': 'English', 'time': '128 min'}
    d = docxwrite()
    save='G://Eclipse workspace/Movie-Aggregator/package1/Result'
    d._write(det,save)
    